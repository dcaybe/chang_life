import docx

try:
    doc = docx.Document('Lộ trình 8 tuần.txt')
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                content.append(' | '.join(row_data))
    
    with open('output.txt', 'w', encoding='utf-8') as f:
        f.write('\n'.join(content))
        
    print("Done writing to output.txt")
except Exception as e:
    print("Error:", e)
